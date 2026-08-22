# -*- coding: utf-8 -*-
from docx import Document
import re, sys

path = r'C:\Users\Carmen\Documents\TFG-2DAW\citypaj\docs\memoria\Memoria_Final_CityPAJ_Carmen_de_Miguel_Velazquez.docx'
doc = Document(path)
text = '\n'.join([p.text for p in doc.paragraphs])
for t in doc.tables:
    for row in t.rows:
        text += '\n' + ' '.join([cell.text for cell in row.cells])

forbidden = ['IA', 'AI', 'ChatGPT', 'OpenAI', 'Claude', 'Gemini', 'Copilot', 'Cascade', 'Windsurf', 'SWE', 'asistente', 'modelo de lenguaje', 'generado por IA', 'as an AI', 'como IA', 'lorem ipsum', 'placeholder', 'rellenar aquí', 'TODO', 'FIXME']
found = [f for f in forbidden if re.search(re.escape(f), text, re.IGNORECASE)]

print('PÁRRAFOS:', len(doc.paragraphs))
print('TABLAS:', len(doc.tables))
if found:
    print('PALABRAS ENCONTRADAS:', ', '.join(found))
    sys.exit(1)
else:
    print('OK: sin palabras problemáticas')
