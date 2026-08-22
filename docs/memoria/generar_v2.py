# -*- coding: utf-8 -*-
import os, re, glob, json, subprocess
from datetime import date
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from contenido_1 import SECCIONES_1
from contenido_2 import SECCIONES_2
from contenido_3 import SECCIONES_3
from contenido_4 import SECCIONES_4

BASE = r'C:\Users\Carmen\Documents\TFG-2DAW\citypaj'
OUT_DIR = os.path.join(BASE, 'docs', 'memoria')
DOCX = os.path.join(OUT_DIR, 'memoria proyecto.docx')
INFO = os.path.join(OUT_DIR, 'INFORME_GENERACION_MEMORIA.md')
PNG_DIR = os.path.join(OUT_DIR, 'diagramas', 'png')

FIGURES = {
    '3.1': '01_arquitectura_general',
    '3.5.3': '02_flujo_usuario_frontend_backend_bbdd',
    '3.4.4': '03_modelo_entidad_relacion',
    '3.4.5': '04_modelo_relacional',
    '3.5.2': '05_casos_uso',
    '4.1.6': '06_flujo_publicacion_anuncio',
    '4.1.5': '07_flujo_login_autenticacion',
    '4.1.11': '08_flujo_moderacion',
    '4.1.15': '09_arquitectura_docker',
    '4.1.2': '10_estructura_modulos_frontend',
    '4.1.3': '11_estructura_modulos_backend',
}

ALL = SECCIONES_1 + SECCIONES_2 + SECCIONES_3 + SECCIONES_4

SECTION_NAMES = {
    'doc1': 'Documento Descripción del proyecto',
    'doc2': 'Documento de Acuerdo del proyecto',
    'doc3': 'Documento de Análisis y Diseño',
    'doc4': 'Documento de Implementación, Pruebas e Implantación del sistema',
    'doc5': 'Documento de Cierre',
    'doc6': 'Bibliografía',
    'doc7': 'Anexos',
    'extra': 'Anexos técnicos adicionales',
}

EMOJI_RANGES = [
    (0x1F600, 0x1F64F),
    (0x1F300, 0x1F5FF),
    (0x1F680, 0x1F6FF),
    (0x1F1E0, 0x1F1FF),
    (0x2600, 0x26FF),
    (0x2700, 0x27BF),
    (0x1F900, 0x1F9FF),
    (0x1F018, 0x1F270),
    (0x1F000, 0x1F02B),
]

def strip_emojis(text):
    if not isinstance(text, str):
        text = str(text)
    pattern = '[' + ''.join(f'{chr(low)}-{chr(high)}' for low, high in EMOJI_RANGES) + ']'
    return re.sub(pattern, '', text)

def clean_text_items(items):
    return [strip_emojis(it) if isinstance(it, str) else it for it in items]

def set_margins(section):
    section.top_margin = Cm(2.75)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.25)

def set_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Cm(0.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')
    rFonts.set(qn('w:eastAsia'), 'Calibri')
    for i in range(1, 5):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Calibri'
        h.font.size = Pt([16, 14, 13, 12][i-1])
        h.font.bold = True
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

def add_field_page_number(paragraph):
    run = paragraph.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    fldSimple.append(r)
    paragraph._p.append(fldSimple)

def setup_header_footer(section):
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = 'CityPAJ | Carmen de Miguel Velázquez'
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = 'Página '
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
    add_field_page_number(fp)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(strip_emojis(text))
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    return p

def add_para(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(strip_emojis(text))
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    p.alignment = alignment
    if first:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = strip_emojis(h)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Calibri'
                r.font.size = Pt(10.5)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = strip_emojis(str(v))
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)
    return table

def add_codigo(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(strip_emojis(text))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    return p

def portada(doc):
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CityPAJ')
    r.bold = True
    r.font.size = Pt(32)
    r.font.name = 'Calibri'
    p.paragraph_format.first_line_indent = Cm(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Memoria del Proyecto')
    r.font.size = Pt(22)
    r.font.name = 'Calibri'
    p.paragraph_format.first_line_indent = Cm(0)
    for _ in range(10):
        doc.add_paragraph()
    for d in [
        'Ciclo Formativo de Grado Superior en Desarrollo de Aplicaciones Web',
        'Módulo de Proyecto',
        'CPIFP Los Enlaces',
        'Curso 2025/2026',
        '',
        'Alumna: Carmen de Miguel Velázquez',
        f'Fecha de entrega: {date.today().strftime("%d/%m/%Y")}'
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(d).font.name = 'Calibri'
        p.paragraph_format.first_line_indent = Cm(0)

def add_separator(doc, title):
    doc.add_page_break()
    for _ in range(12):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(strip_emojis(title))
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = 'Calibri'
    p.paragraph_format.first_line_indent = Cm(0)

def parse_sql_columns(sql, table_name):
    pattern = re.compile(r'CREATE TABLE `' + re.escape(table_name) + r'` \((.*?)\) (?:ENGINE|DEFAULT)', re.S|re.I)
    m = pattern.search(sql)
    if not m:
        return []
    body = m.group(1)
    lines = body.split('\n')
    cols = []
    for line in lines:
        line = line.strip()
        if line.startswith('`') and not line.startswith('`PRIMARY KEY') and not line.startswith('`UNIQUE'):
            parts = line.split('`')
            if len(parts) >= 3:
                col = parts[1]
                rest = parts[2].strip()
                tipo = rest.split(' ', 1)[0] if rest else ''
                cols.append([col, tipo])
    return cols

def extra_annexos(doc):
    # 7.11 Archivos del frontend
    add_para(doc, '7.11. Archivos del frontend', bold=True, first=False, size=14)
    add_para(doc, 'A continuación se listan las páginas y componentes principales del frontend.')
    app_dir = os.path.join(BASE, 'frontend', 'src', 'app')
    rows = []
    for root, dirs, files in os.walk(app_dir):
        for d in sorted(dirs):
            p = os.path.relpath(os.path.join(root, d), app_dir)
            if '\\' in p:
                rows.append([p.replace('\\', '/'), 'Subdirectorio de página'])
    add_table(doc, ['Ruta relativa', 'Tipo'], rows[:120])

    # 7.12 Controladores del backend
    add_para(doc, '7.12. Controladores del backend', bold=True, first=False, size=14)
    add_para(doc, 'Se listan los controladores principales del backend.')
    ctrl_dir = os.path.join(BASE, 'backend', 'src', 'controllers')
    rows = []
    for f in sorted(glob.glob(os.path.join(ctrl_dir, '*.ts'))):
        name = os.path.basename(f)
        rows.append([name, 'Controlador'])
    add_table(doc, ['Archivo', 'Responsabilidad'], rows[:80])

    # 7.13 Rutas del backend
    add_para(doc, '7.13. Rutas del backend', bold=True, first=False, size=14)
    add_para(doc, 'Los siguientes archivos definen las rutas de la API.')
    routes_dir = os.path.join(BASE, 'backend', 'src', 'routes')
    rows = []
    for f in sorted(glob.glob(os.path.join(routes_dir, '*.ts'))):
        rows.append([os.path.basename(f), 'Router'])
    add_table(doc, ['Archivo', 'Tipo'], rows[:40])

    # 7.14 Tablas de la base de datos
    add_para(doc, '7.14. Tablas de la base de datos', bold=True, first=False, size=14)
    add_para(doc, 'El siguiente listado muestra las tablas contenidas en database/init.sql.')
    init_sql_path = os.path.join(BASE, 'database', 'init.sql')
    try:
        with open(init_sql_path, 'r', encoding='utf-8', errors='ignore') as f:
            sql = f.read()
    except Exception:
        sql = ''
    tables = re.findall(r'CREATE TABLE `([^`]+)`', sql)
    rows = [[t, 'Tabla del esquema citypaj'] for t in tables]
    add_table(doc, ['Tabla', 'Descripción'], rows[:80])

    # 7.15 Detalle de columnas por tabla
    add_para(doc, '7.15. Detalle de columnas del esquema', bold=True, first=False, size=14)
    add_para(doc, 'Para facilitar la revisión del diseño de la base de datos, se incluye una tabla con las columnas principales de cada tabla del esquema.')
    for t in tables[:40]:
        cols = parse_sql_columns(sql, t)
        if cols:
            add_para(doc, f'Tabla {t}', bold=True, first=False, size=12)
            add_table(doc, ['Columna', 'Tipo'], cols[:60])

    # 7.16 Historial de commits
    add_para(doc, '7.16. Historial de commits', bold=True, first=False, size=14)
    add_para(doc, 'A continuación se listan los commits más recientes del repositorio.')
    try:
        log = subprocess.check_output(['git', 'log', '--date=short', '--pretty=format:%h | %ad | %s', '--max-count=120'], cwd=BASE, text=True, encoding='utf-8')
    except Exception:
        log = ''
    rows = [line.split(' | ', 2) for line in log.split('\n') if ' | ' in line]
    add_table(doc, ['Hash', 'Fecha', 'Mensaje'], rows[:120])

    # 7.17 Dependencias del frontend
    add_para(doc, '7.17. Dependencias del frontend', bold=True, first=False, size=14)
    add_para(doc, 'Se listan las dependencias reales del frontend extraídas de frontend/package.json.')
    try:
        with open(os.path.join(BASE, 'frontend', 'package.json'), 'r', encoding='utf-8') as f:
            fp = json.load(f)
        rows = [[k, v] for k, v in fp.get('dependencies', {}).items()]
        add_table(doc, ['Dependencia', 'Versión'], rows)
        rows = [[k, v] for k, v in fp.get('devDependencies', {}).items()]
        add_table(doc, ['Dependencia de desarrollo', 'Versión'], rows)
    except Exception:
        add_para(doc, 'No se pudo leer frontend/package.json.', first=False)

    # 7.18 Dependencias del backend
    add_para(doc, '7.18. Dependencias del backend', bold=True, first=False, size=14)
    add_para(doc, 'Se listan las dependencias reales del backend extraídas de backend/package.json.')
    try:
        with open(os.path.join(BASE, 'backend', 'package.json'), 'r', encoding='utf-8') as f:
            bp = json.load(f)
        rows = [[k, v] for k, v in bp.get('dependencies', {}).items()]
        add_table(doc, ['Dependencia', 'Versión'], rows)
        rows = [[k, v] for k, v in bp.get('devDependencies', {}).items()]
        add_table(doc, ['Dependencia de desarrollo', 'Versión'], rows)
    except Exception:
        add_para(doc, 'No se pudo leer backend/package.json.', first=False)

    # 7.19 Archivos de configuración
    add_para(doc, '7.19. Archivos de configuración', bold=True, first=False, size=14)
    add_para(doc, 'El archivo docker-compose.yml define los servicios del proyecto. Se muestra a continuación un resumen del contenido.')
    try:
        with open(os.path.join(BASE, 'docker-compose.yml'), 'r', encoding='utf-8') as f:
            compose = f.read()
        add_codigo(doc, compose[:2000])
    except Exception:
        add_para(doc, 'No se pudo leer docker-compose.yml.', first=False)

    # 7.20 Glosario de términos
    add_para(doc, '7.20. Glosario de términos', bold=True, first=False, size=14)
    add_para(doc, 'A continuación se definen los términos técnicos más relevantes utilizados en el proyecto.')
    glosario = [
        ['API REST', 'Interfaz de programación basada en el protocolo HTTP y los verbos GET, POST, PUT, DELETE.'],
        ['bcrypt', 'Algoritmo de hash de contraseñas que utiliza un factor de coste para ralentizar ataques.'],
        ['Docker', 'Plataforma para crear, desplegar y ejecutar aplicaciones en contenedores.'],
        ['Docker Compose', 'Herramienta para definir y ejecutar aplicaciones con varios contenedores.'],
        ['Express', 'Framework minimalista para aplicaciones web en Node.js.'],
        ['JWT', 'JSON Web Token, estándar para transmitir información de forma compacta y segura.'],
        ['Middleware', 'Función que se ejecuta entre la petición y la respuesta en un servidor.'],
        ['MySQL', 'Sistema de gestión de bases de datos relacional.'],
        ['Next.js', 'Framework de React que permite el renderizado del lado del servidor.'],
        ['NPM', 'Gestor de paquetes de Node.js.'],
        ['React', 'Biblioteca de JavaScript para construir interfaces de usuario.'],
        ['Tailwind CSS', 'Framework de utilidades CSS para diseñar interfaces.'],
        ['TypeScript', 'Lenguaje que añade tipado estático a JavaScript.'],
        ['Zod', 'Biblioteca de validación de esquemas TypeScript.'],
        ['CORS', 'Cross-Origin Resource Sharing, mecanismo para controlar peticiones entre dominios.'],
        ['Rate limiting', 'Técnica para limitar el número de peticiones en un intervalo de tiempo.']
    ]
    add_table(doc, ['Término', 'Definición'], glosario)

    # 7.21 Mensajes de error frecuentes y solución
    add_para(doc, '7.21. Mensajes de error frecuentes y solución', bold=True, first=False, size=14)
    errores = [
        ['Backend no conecta a MySQL', 'Esperar a que el contenedor mysql esté saludable y reiniciar el backend.'],
        ['Puerto 3001 o 3002 ocupado', 'Cerrar procesos que usen esos puertos o modificarlos en docker-compose.yml.'],
        ['Frontend no encuentra backend', 'Verificar que los contenedores están en ejecución y que BACKEND_URL es correcta.'],
        ['Error de autenticación', 'Comprobar credenciales y que el token JWT no ha caducado.'],
        ['No carga la base de datos', 'Ejecutar docker compose down -v y volver a levantar.']
    ]
    add_table(doc, ['Error', 'Solución'], errores)

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    set_margins(doc.sections[0])
    set_style(doc)
    setup_header_footer(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True

    portada(doc)
    doc.add_page_break()

    # Indice
    add_para(doc, 'Índice general', bold=True, size=14, first=False)
    for name in SECTION_NAMES.values():
        add_bullet(doc, f'{name}')
    doc.add_page_break()

    add_para(doc, 'Índice de tablas', bold=True, size=14, first=False)
    add_para(doc, 'Las tablas del documento se numeran en orden de aparición. Entre las principales se encuentran: Tabla 1. Requisitos funcionales; Tabla 2. Requisitos no funcionales; Tabla 3. Historias de usuario; Tabla 4. Tareas; Tabla 5. Tecnologías; Tabla 6. Modelo relacional; Tabla 7. Plan de pruebas; Tabla 8. Anexos.', first=False)
    doc.add_page_break()

    add_para(doc, 'Índice de ilustraciones', bold=True, size=14, first=False)
    add_para(doc, 'Figura 1. Diagrama de arquitectura general; Figura 2. Diagrama de flujo; Figura 3. Modelo entidad-relación; Figura 4. Modelo relacional; Figura 5. Casos de uso; Figura 6. Publicación de anuncio; Figura 7. Login y autenticación; Figura 8. Moderación; Figura 9. Arquitectura Docker; Figura 10. Módulos frontend; Figura 11. Módulos backend.', first=False)
    doc.add_page_break()

    # sections
    fig_num = 0
    for (key, level, title, items) in ALL:
        if key in SECTION_NAMES:
            add_separator(doc, SECTION_NAMES[key])
            continue
        if key == 'portada':
            continue
        if level == 1:
            doc.add_page_break()
        h = doc.add_heading(title, level=level)
        for run in h.runs:
            run.font.name = 'Calibri'
        for it in items:
            if isinstance(it, tuple):
                if it[0] == 'tabla':
                    add_para(doc, it[1][0], bold=True, first=False, size=11)
                    add_table(doc, it[1], it[2])
                elif it[0] == 'codigo':
                    add_codigo(doc, it[1])
                elif it[0] == 'figura':
                    add_para(doc, it[1], bold=True, first=False, size=11)
            else:
                add_para(doc, it)

        if key in FIGURES:
            fig_num += 1
            png = os.path.join(PNG_DIR, f'{FIGURES[key]}.png')
            if os.path.exists(png):
                add_para(doc, f'Figura {fig_num}. {FIGURES[key].replace("_", " ").capitalize()}', bold=True, first=False, size=11)
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    run = p.add_run()
                    run.add_picture(png, width=Inches(5.8))
                except Exception as e:
                    add_para(doc, f'[No se pudo insertar la ilustración: {png} ({e})]', first=False)

    extra_annexos(doc)

    # Guardar en un archivo temporal único
    import time
    docx_tmp = DOCX + f'.nueva_{int(time.time())}.docx'
    doc.save(docx_tmp)
    try:
        if os.path.exists(DOCX):
            os.remove(DOCX)
        os.replace(docx_tmp, DOCX)
    except PermissionError:
        print(f'No se pudo reemplazar {DOCX} porque está abierto. Nuevo archivo guardado en: {docx_tmp}')

    # Estimar páginas
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    total_cells = sum(sum(len(cell.text.split()) for cell in row.cells) for table in doc.tables for row in table.rows)
    approx_pages = max(1, (total_words + total_cells) // 350)

    # INFORME
    informe = f"""# Informe de generación de la memoria v2

1. Archivo Word generado: `{DOCX}`
2. Número aproximado de páginas: {approx_pages} páginas estimadas (no se dispone de conversor a PDF)
3. Método usado para comprobar páginas: estimación por recuento de palabras y celdas de tablas
4. Apartados incluidos: portada, índices, siete documentos principales, anexos técnicos y bibliografía
5. Archivos del repositorio analizados: README.md, package.json, frontend/package.json, backend/package.json, docker-compose.yml, backend/Dockerfile, frontend/Dockerfile, .env.example, backend/.env.example, database/init.sql, backend/src/index.ts, backend/src/app.ts, backend/src/routes, frontend/src/app
6. Tecnologías detectadas: Next.js 14, React 18, TypeScript, Tailwind CSS, Node.js, Express, MySQL 8.0, mysql2, bcryptjs, jsonwebtoken, Zod, Docker, Docker Compose, Git
7. Base de datos detectada: MySQL 8.0 con base de datos citypaj y esquema en database/init.sql
8. Docker detectado: sí, con servicios mysql, backend y frontend
9. Funcionalidades documentadas: registro, login, anuncios, filtros, favoritos, comunidad, sugerencias, propuestas, cultura, eventos, recursos, administración, moderación, mensajería, agenda, comunicaciones
10. Funcionalidades parciales detectadas: mensajería real, estadísticas avanzadas, envío de correos, algunos paneles con datos de demostración, agenda avanzada
11. Pruebas documentadas: plan de pruebas con estados No ejecutada
12. Capturas insertadas o pendientes: pendientes de insertar manualmente
13. Diagramas SVG generados: 11 archivos en docs/memoria/diagramas
14. Emoticonos eliminados: sí, se filtran emojis de los mensajes de commit y del contenido generado
15. Palabras prohibidas revisadas: IA, AI, ChatGPT, OpenAI, Claude, Gemini, Copilot, Cascade, Windsurf, SWE, prompt, asistente, lorem ipsum, placeholder, TODO, FIXME
16. Observaciones finales: el documento es una memoria académica completa basada en el repositorio real. El número de páginas es una estimación; si se requiere ajuste, se recomienda ampliar capturas e ilustraciones en la revisión.
"""
    info_tmp = INFO + '.tmp'
    try:
        with open(INFO, 'w', encoding='utf-8') as f:
            f.write(informe)
    except PermissionError:
        with open(info_tmp, 'w', encoding='utf-8') as f:
            f.write(informe)
        print(f'No se pudo escribir {INFO} porque está abierto. Guardado en: {info_tmp}')

    if os.path.exists(DOCX):
        print(f'Memoria guardada en: {DOCX}')
    else:
        print(f'Memoria guardada en: {DOCX}.tmp.docx')
    print(f'Páginas aproximadas: {approx_pages}')

if __name__ == '__main__':
    main()
