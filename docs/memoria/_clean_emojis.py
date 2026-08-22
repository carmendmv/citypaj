# -*- coding: utf-8 -*-
import os
import re
from docx import Document

BASE = r'C:\Users\Carmen\Documents\TFG-2DAW\citypaj\docs\memoria'
SRC = os.path.join(BASE, 'Memoria_Final_CityPAJ_Carmen_de_Miguel_Velazquez.docx.tmp.docx')
DST = os.path.join(BASE, 'Memoria_Final_CityPAJ_Carmen_de_Miguel_Velazquez.docx')

EMOJI_RANGES = [
    (0x1F600, 0x1F64F),
    (0x1F300, 0x1F5FF),
    (0x1F680, 0x1F6FF),
    (0x1F1E0, 0x1F1FF),
    (0x2600, 0x26FF),
    (0x2700, 0x27BF),
    (0x1F900, 0x1F9FF),
    (0x1F018, 0x1F270),
    (0x1F000, 0x1F02B),
]

def strip_emojis(text):
    if not isinstance(text, str):
        text = str(text)
    pattern = '[' + ''.join(f'{chr(low)}-{chr(high)}' for low, high in EMOJI_RANGES) + ']'
    return re.sub(pattern, '', text)

def clean_doc(path):
    doc = Document(path)
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = strip_emojis(run.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = strip_emojis(run.text)
    # Guardar en destino si es posible, si no en temporal
    try:
        doc.save(DST)
        print(f'Guardado en: {DST}')
    except PermissionError:
        tmp = DST + '.tmp.docx'
        doc.save(tmp)
        print(f'No se pudo reemplazar {DST}. Guardado en: {tmp}')

if __name__ == '__main__':
    if os.path.exists(SRC):
        clean_doc(SRC)
    else:
        print(f'No se encuentra {SRC}')
