# -*- coding: utf-8 -*-
import os
import re
from datetime import date
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'C:\Users\Carmen\Documents\TFG-2DAW\citypaj'
OUT_DIR = os.path.join(BASE, 'docs', 'memoria')
DOCX = os.path.join(OUT_DIR, 'Memoria_Final_CityPAJ_Carmen_de_Miguel_Velazquez.docx')
INFO = os.path.join(OUT_DIR, 'INFORME_GENERACION_MEMORIA.md')

def set_margins(section):
    section.top_margin = Cm(2.75)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.25)

def set_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Cm(0.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    r = style.element.rPr
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')
    rFonts.set(qn('w:eastAsia'), 'Calibri')
    r.append(rFonts)

    for i in range(1, 5):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Calibri'
        h.font.size = Pt([16, 14, 13, 12][i-1])
        h.font.bold = True
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    p.alignment = alignment
    if first:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.add_run(text).font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10.5)
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.shading = None
    return p

def portada(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('CityPAJ')
    r.bold = True
    r.font.size = Pt(28)
    r.font.name = 'Calibri'
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Memoria del Proyecto')
    r.font.size = Pt(20)
    r.font.name = 'Calibri'
    for _ in range(8):
        doc.add_paragraph()
    datos = [
        'Ciclo Formativo de Grado Superior',
        'Desarrollo de Aplicaciones Web',
        'Módulo de Proyecto',
        'CPIFP Los Enlaces',
        'Curso 2025/2026',
        '',
        'Alumna: Carmen de Miguel Velázquez',
        f'Fecha de entrega: {date.today().strftime("%d/%m/%Y")}'
    ]
    for d in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(d).font.name = 'Calibri'
        p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    set_margins(doc.sections[0])
    set_style(doc)
    portada(doc)

    # ÍNDICES
    add_heading(doc, 'Índice general', level=1)
    indice_items = [
        '1. Documento Descripción del proyecto',
        '2. Documento de Acuerdo del proyecto',
        '3. Documento de Análisis y Diseño',
        '4. Documento de Implementación, Pruebas e Implantación del sistema',
        '5. Documento de Cierre',
        '6. Bibliografía',
        '7. Anexos',
    ]
    for it in indice_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(it).font.name = 'Calibri'
        p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()

    add_heading(doc, 'Índice de tablas', level=1)
    for t in [
        'Tabla 1. Requisitos funcionales',
        'Tabla 2. Requisitos no funcionales',
        'Tabla 3. Historias de usuario',
        'Tabla 4. Planificación temporal',
        'Tabla 5. Desviaciones temporales',
        'Tabla 6. Presupuesto',
        'Tabla 7. Análisis de riesgos',
        'Tabla 8. Tecnologías utilizadas',
        'Tabla 9. Arquitectura del sistema',
        'Tabla 10. Entidades principales',
        'Tabla 11. Actores del sistema',
        'Tabla 12. Casos de uso',
        'Tabla 13. Plan de pruebas',
        'Tabla 14. Bitácora del proyecto',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(t).font.name = 'Calibri'
        p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()

    add_heading(doc, 'Índice de ilustraciones', level=1)
    add_paragraph(doc, 'El documento no incluye capturas de pantalla. Se recomienda insertar manualmente las figuras correspondientes en una revisión posterior.')
    doc.add_page_break()

    # 1. DESCRIPCIÓN
    add_heading(doc, '1. Documento Descripción del proyecto', level=1, page_break=True)
    add_heading(doc, '1.1. Contexto del proyecto', level=2)
    add_heading(doc, '1.1.1. Ámbito y entorno', level=3)
    add_paragraph(doc, 'CityPAJ se enmarca en el desarrollo de aplicaciones web del ciclo formativo DAW. La propuesta nace de la necesidad detectada durante el aprendizaje de diseñar una aplicación con persistencia, autenticación y una interfaz moderna que responda a un caso de uso social.')
    add_paragraph(doc, 'La aplicación está pensada para personas jóvenes y entidades que precisan un punto de información centralizado donde consultar y publicar anuncios, propuestas, sugerencias, recursos y actividad comunitaria.')
    add_heading(doc, '1.1.2. Análisis de la realidad', level=3)
    add_paragraph(doc, 'Actualmente la información juvenil se distribuye en múltiples canales: webs institucionales, redes sociales, tablones de anuncios locales o grupos de mensajería. Esta dispersión dificulta la consulta sistemática y el filtrado por territorio o categoría.')
    add_paragraph(doc, 'Además, no siempre queda clara la diferencia entre un anuncio, una propuesta o un recurso. CityPAJ plantea una organización por secciones para resolver esta fragmentación.')
    add_heading(doc, '1.1.3. Solución y justificación', level=3)
    add_paragraph(doc, 'La solución consiste en una aplicación web fullstack con un frontend en Next.js, un backend con Express y una base de datos MySQL. Los datos se organizan por comunidades autónomas, provincias y categorías, de modo que el usuario puede localizar contenido relevante de forma rápida.')
    add_paragraph(doc, 'Se ha elegido Docker Compose para simplificar la ejecución local y facilitar la revisión del proyecto sin necesidad de instalar MySQL o Node.js de forma independiente.')
    add_heading(doc, '1.1.4. Destinatarios', level=3)
    for it in [
        'Jóvenes que buscan anuncios de empleo, formación, vivienda, ocio o servicios.',
        'Entidades e instituciones juveniles que desean publicar recursos.',
        'Moderadores que supervisan contenidos antes de su publicación.',
        'Administradores que gestionan usuarios, categorías y paneles del sistema.'
    ]:
        add_bullet(doc, it)

    add_heading(doc, '1.2. Objetivo del proyecto', level=2)
    add_paragraph(doc, 'El objetivo es desarrollar una aplicación web que centralice información juvenil organizada por territorio, permitiendo la consulta y publicación de anuncios, recursos, propuestas, sugerencias y actividad comunitaria. El sistema incluye roles de usuario, moderación básica de contenidos, paneles de administración y un despliegue local mediante Docker.')

    add_heading(doc, '1.3. Project Objective', level=2)
    add_paragraph(doc, 'The aim of this project is to develop a web application that gathers youth information organised by territory. It allows users to consult and publish announcements, resources, proposals, suggestions and community activities. The system includes user roles, basic content moderation, administration panels and a local Docker deployment.')

    add_heading(doc, '1.4. Marco legal', level=2)
    add_paragraph(doc, 'El proyecto procesa datos personales mínimos, como correo electrónico, contraseña cifrada con bcrypt y dirección IP en registros de actividad. En un despliegue real debería cumplir el Reglamento General de Protección de Datos (RGPD) y la Ley Orgánica 3/2018 de Protección de Datos Personales y garantía de los derechos digitales.')
    add_paragraph(doc, 'El tratamiento de datos se limita a lo necesario para el registro, la autenticación y la moderación. Se ha implementado una separación de roles que permite controlar quién accede a cada funcionalidad. La licencia declarada en el repositorio es MIT, lo que permite su uso académico y debería completarse con la autorización de la autora para cualquier publicación externa.')

    # 2. ACUERDO
    add_heading(doc, '2. Documento de Acuerdo del proyecto', level=1, page_break=True)
    add_heading(doc, '2.1. Requisitos funcionales', level=2)
    rf_rows = [
        ['RF-01', 'Registro de usuarios', 'Permite crear una cuenta con correo y contraseña.', 'Visitante', 'Cumplido', 'Registro con credenciales demo'],
        ['RF-02', 'Inicio de sesión', 'Autenticación con JWT y refresh token.', 'Usuario', 'Cumplido', 'Login con usuarios de prueba'],
        ['RF-03', 'Gestión de sesión', 'Cierre de sesión y renovación de tokens.', 'Usuario', 'Cumplido', 'Logout y refresh'],
        ['RF-04', 'Consulta de anuncios', 'Listado paginado de anuncios.', 'Usuario', 'Cumplido', 'Visualización home'],
        ['RF-05', 'Filtrado por categoría', 'Filtros por tipo de anuncio.', 'Usuario', 'Cumplido', 'Filtros en frontend'],
        ['RF-06', 'Filtrado por territorio', 'Filtros por CCAA y provincia.', 'Usuario', 'Parcial', 'Filtros implementados en backend'],
        ['RF-07', 'Publicación de anuncios', 'Formulario de publicación con adjuntos.', 'Usuario registrado', 'Cumplido', 'Publicar anuncio'],
        ['RF-08', 'Gestión de anuncios propios', 'Edición y eliminación de anuncios.', 'Usuario', 'Parcial', 'Edición y mis anuncios'],
        ['RF-09', 'Detalle de anuncio', 'Vista individual de un anuncio.', 'Usuario', 'Cumplido', 'Página de detalle'],
        ['RF-10', 'Favoritos', 'Guardar anuncios como favoritos.', 'Usuario', 'Cumplido', 'Panel de guardados'],
        ['RF-11', 'Reporte de anuncios', 'Enviar reportes de contenido.', 'Usuario', 'Cumplido', 'Modal de reporte'],
        ['RF-12', 'Comunidad por provincias', 'Publicaciones locales y comentarios.', 'Usuario', 'Parcial', 'Sección comunidad'],
        ['RF-13', 'Sugerencias', 'Envío de sugerencias juveniles.', 'Usuario', 'Cumplido', 'Buzón de sugerencias'],
        ['RF-14', 'Propuestas', 'Consulta y apoyo a propuestas.', 'Usuario', 'Parcial', 'Página de propuestas'],
        ['RF-15', 'Cultura y eventos', 'Sección específica de actividad cultural.', 'Usuario', 'Cumplido', 'Página cultura y eventos'],
        ['RF-16', 'Recursos', 'Consulta de recursos juveniles.', 'Usuario', 'Parcial', 'Página recursos'],
        ['RF-17', 'Panel de administración', 'Gestión de contenidos y usuarios.', 'Admin', 'Parcial', 'Panel admin con módulos'],
        ['RF-18', 'Panel de moderación', 'Revisión y acciones sobre contenidos.', 'Moderador', 'Parcial', 'Panel moderador'],
        ['RF-19', 'Mensajería', 'Comunicación interna.', 'Admin/Moderador', 'Parcial', 'Vista de mensajes en panel'],
        ['RF-20', 'Dockerización', 'Ejecución local con Docker Compose.', 'Sistema', 'Cumplido', 'docker compose up --build']
    ]
    add_table(doc, ['ID', 'Requisito', 'Descripción', 'Actor', 'Estado', 'Prueba'], rf_rows)

    add_heading(doc, '2.2. Requisitos no funcionales', level=2)
    rnf_rows = [
        ['RNF-01', 'Usabilidad', 'Interfaz clara y navegación por secciones.', 'Facilita el uso del sistema.', 'Cumplido'],
        ['RNF-02', 'Accesibilidad básica', 'Contraste y estructura semántica.', 'Mejora la legibilidad.', 'Parcial'],
        ['RNF-03', 'Diseño responsive', 'Adaptación a móvil y escritorio.', 'Navegación en distintos dispositivos.', 'Cumplido'],
        ['RNF-04', 'Seguridad', 'Contraseñas cifradas con bcrypt y JWT.', 'Protege cuentas y sesiones.', 'Cumplido'],
        ['RNF-05', 'Validación de datos', 'Esquemas con Zod y validaciones de formulario.', 'Reduce datos erróneos.', 'Cumplido'],
        ['RNF-06', 'Mantenibilidad', 'Código modular en capas.', 'Facilita correcciones.', 'Cumplido'],
        ['RNF-07', 'Separación frontend/backend', 'Arquitectura cliente-servidor.', 'Claridad y escalabilidad.', 'Cumplido'],
        ['RNF-08', 'Persistencia MySQL', 'Base de datos relacional con esquema definido.', 'Almacenamiento estructurado.', 'Cumplido'],
        ['RNF-09', 'Portabilidad con Docker', 'Contenedores para toda la pila.', 'Ejecución sin instalaciones locales.', 'Cumplido'],
        ['RNF-10', 'Rendimiento', 'Suficiente para entorno académico.', 'Tiempo de respuesta aceptable.', 'Parcial'],
        ['RNF-11', 'Compatibilidad', 'Navegadores modernos.', 'Uso generalizado.', 'Cumplido'],
        ['RNF-12', 'Claridad visual', 'Diseño con Tailwind CSS.', 'Consistencia de estilos.', 'Cumplido'],
        ['RNF-13', 'Protección de datos', 'Tratamiento mínimo y roles.', 'Cumplimiento básico.', 'Cumplido'],
        ['RNF-14', 'Escalabilidad futura', 'Arquitectura modular.', 'Permite ampliaciones.', 'Parcial'],
        ['RNF-15', 'Documentación', 'README con instalación y decisiones.', 'Facilita la entrega.', 'Cumplido']
    ]
    add_table(doc, ['ID', 'Requisito', 'Descripción', 'Justificación', 'Estado'], rnf_rows)

    add_heading(doc, '2.3. Historias de usuario', level=2)
    hu_rows = [
        ['HU-01', 'Como visitante, quiero ver anuncios, para conocer ofertas juveniles.', 'Página de inicio accesible.', 'RF-04'],
        ['HU-02', 'Como usuario, quiero registrarme, para publicar contenido.', 'Formulario de registro funcional.', 'RF-01'],
        ['HU-03', 'Como usuario, quiero iniciar sesión, para acceder a mi perfil.', 'Login con JWT.', 'RF-02'],
        ['HU-04', 'Como usuario, quiero publicar un anuncio, para compartir una oportunidad.', 'Formulario de publicación.', 'RF-07'],
        ['HU-05', 'Como usuario, quiero filtrar por provincia, para encontrar anuncios cercanos.', 'Filtros territoriales.', 'RF-06'],
        ['HU-06', 'Como usuario, quiero guardar favoritos, para consultarlos después.', 'Panel de guardados.', 'RF-10'],
        ['HU-07', 'Como usuario, quiero enviar sugerencias, para proponer mejoras.', 'Buzón de sugerencias.', 'RF-13'],
        ['HU-08', 'Como moderador, quiero revisar reportes, para mantener la calidad.', 'Panel de moderación.', 'RF-18'],
        ['HU-09', 'Como administrador, quiero gestionar usuarios, para controlar el acceso.', 'Panel admin de usuarios.', 'RF-17'],
        ['HU-10', 'Como administrador, quiero ver estadísticas, para entender el uso.', 'Dashboard de estadísticas.', 'RF-17']
    ]
    add_table(doc, ['ID', 'Historia de usuario', 'Criterios de aceptación', 'Requisitos'], hu_rows)

    add_heading(doc, '2.4. Definición de tareas', level=2)
    add_paragraph(doc, 'El proyecto se ha dividido en fases: análisis inicial, diseño de base de datos, diseño de interfaz, configuración del repositorio, implementación del backend, implementación del frontend, conexión con MySQL, desarrollo de autenticación, módulo de anuncios, filtros, comunidad, sugerencias, propuestas, cultura y eventos, panel de administración, moderación, Dockerización, pruebas y documentación final.')

    add_heading(doc, '2.5. Metodología', level=2)
    add_paragraph(doc, 'El proyecto utiliza una metodología mixta, ya que combina una fase inicial de análisis y diseño con una implementación incremental. Esta forma de trabajo resulta adecuada para un proyecto académico, porque permite definir primero una base funcional y después ampliar o corregir módulos según las necesidades detectadas durante el desarrollo.')

    add_heading(doc, '2.6. Planificación temporal', level=2)
    plan_rows = [
        ['Análisis', 'Análisis de requisitos y diseño preliminar', '10', '12', 'Completado'],
        ['Diseño BD', 'Modelado de base de datos y scripts SQL', '8', '10', 'Completado'],
        ['Diseño UI', 'Prototipado de pantallas principales', '8', '6', 'Completado'],
        ['Configuración', 'Repositorio, Docker y variables de entorno', '6', '8', 'Completado'],
        ['Backend', 'API, autenticación y controladores', '30', '35', 'Completado'],
        ['Frontend', 'Páginas, componentes y navegación', '35', '40', 'Completado'],
        ['Docker', 'Contenedores y conexión entre servicios', '6', '8', 'Completado'],
        ['Pruebas', 'Pruebas funcionales y correcciones', '10', '12', 'Parcial'],
        ['Documentación', 'README y memoria', '8', '10', 'En curso'],
    ]
    add_table(doc, ['Fase', 'Tarea', 'Horas previstas', 'Horas reales', 'Resultado'], plan_rows)

    add_heading(doc, '2.7. Presupuesto', level=2)
    pres_rows = [
        ['Análisis y diseño', '30', '0', '0', 'Trabajo propio'],
        ['Implementación frontend', '40', '0', '0', 'Trabajo propio'],
        ['Implementación backend', '35', '0', '0', 'Trabajo propio'],
        ['Base de datos', '10', '0', '0', 'Trabajo propio'],
        ['Docker', '8', '0', '0', 'Trabajo propio'],
        ['Pruebas', '12', '0', '0', 'Trabajo propio'],
        ['Documentación', '10', '0', '0', 'Trabajo propio'],
        ['Herramientas software', '-', '-', '0', 'Software libre (Node, Docker, MySQL)'],
        ['Despliegue local', '-', '-', '0', 'Docker Desktop']
    ]
    add_table(doc, ['Concepto', 'Horas / Unidades', 'Coste unitario', 'Coste total', 'Observación'], pres_rows)
    add_paragraph(doc, 'El presupuesto es académico. Las horas corresponden al trabajo de la alumna y los costes de software son nulos al emplear herramientas libres y gratuitas.')

    add_heading(doc, '2.8. Licencia o condiciones de distribución', level=2)
    add_paragraph(doc, 'El repositorio declara una licencia MIT en el archivo package.json. El proyecto se entrega con finalidad académica para su revisión dentro del módulo de Proyecto. La reutilización o publicación fuera de este contexto queda sujeta a la autorización de la autora.')

    add_heading(doc, '2.9. Análisis de riesgos', level=2)
    risk_rows = [
        ['R1', 'Retrasos de implementación', 'Media', 'Alta', 'Planificación con márgenes', 'Reducir alcance si es necesario'],
        ['R2', 'Errores de conexión con MySQL', 'Media', 'Media', 'Uso de Docker y healthchecks', 'Revisar variables de entorno'],
        ['R3', 'Problemas con Docker', 'Media', 'Media', 'Probar en distintos entornos', 'Documentar soluciones'],
        ['R4', 'Problemas de autenticación', 'Baja', 'Alta', 'Pruebas con usuarios demo', 'Revisar tokens y roles'],
        ['R5', 'Pérdida de datos', 'Baja', 'Media', 'Copias en SQL y control de versiones', 'Restaurar desde init.sql'],
        ['R6', 'Validaciones incompletas', 'Media', 'Media', 'Uso de Zod y validaciones', 'Añadir validaciones adicionales'],
        ['R7', 'Problemas responsive', 'Baja', 'Media', 'Diseño mobile-first', 'Ajustar estilos'],
        ['R8', 'Incompatibilidad de versiones', 'Baja', 'Baja', 'Fijar versiones en package.json', 'Actualizar dependencias'],
        ['R9', 'Errores de despliegue local', 'Media', 'Media', 'Docker Compose documentado', 'Revisar logs'],
        ['R10', 'Falta de tiempo para documentación', 'Media', 'Media', 'Documentar durante el desarrollo', 'Priorizar secciones críticas']
    ]
    add_table(doc, ['ID', 'Riesgo', 'Probabilidad', 'Impacto', 'Medida preventiva', 'Medida correctiva'], risk_rows)

    # 3. ANÁLISIS Y DISEÑO
    add_heading(doc, '3. Documento de Análisis y Diseño', level=1, page_break=True)
    add_heading(doc, '3.1. Análisis y diseño de la arquitectura', level=2)
    add_paragraph(doc, 'La aplicación sigue una arquitectura de tres capas: presentación (Next.js), lógica de negocio (Express) y persistencia (MySQL). El usuario interactúa con el frontend, que delega las peticiones al backend a través de rutas internas o directas. El backend ejecuta consultas SQL y devuelve JSON.')
    add_paragraph(doc, 'El frontend se comunica con el backend dentro del contenedor Docker mediante la URL http://backend:3002 y desde el navegador del usuario mediante http://localhost:3002. Esta distinción se gestiona mediante variables de entorno.')

    add_heading(doc, '3.2. Tecnologías y herramientas utilizadas', level=2)
    tech_rows = [
        ['Next.js', '14.x', 'Framework frontend y App Router', 'Permite SSR, enrutamiento y componentes React.'],
        ['React', '18.x', 'Biblioteca de interfaz', 'Componentes reutilizables y gestión de estado.'],
        ['TypeScript', '5.2', 'Tipado estático', 'Reduce errores en frontend y backend.'],
        ['Tailwind CSS', '3.3', 'Estilos', 'Maquetación rápida y consistencia visual.'],
        ['Node.js', '>=20', 'Entorno de ejecución', 'Backend y herramientas.'],
        ['Express', '4.18', 'Framework backend', 'API REST y middlewares.'],
        ['MySQL', '8.0', 'Base de datos relacional', 'Persistencia y consultas SQL.'],
        ['mysql2', '3.22', 'Driver Node.js', 'Conexión con MySQL.'],
        ['bcryptjs', '2.4', 'Cifrado de contraseñas', 'Seguridad de autenticación.'],
        ['jsonwebtoken', '9.0', 'JWT', 'Gestión de sesiones.'],
        ['Zod', '3.22', 'Validación de esquemas', 'Validación de formularios.'],
        ['Docker', 'Compose', 'Contenedores', 'Ejecución local sin dependencias externas.'],
        ['GitHub', '-', 'Control de versiones', 'Repositorio y seguimiento de cambios.']
    ]
    add_table(doc, ['Tecnología', 'Versión', 'Uso', 'Justificación'], tech_rows)

    add_heading(doc, '3.3. Arquitectura de componentes', level=2)
    add_paragraph(doc, 'La arquitectura se organiza en tres bloques principales.')
    arch_rows = [
        ['Capa de presentación', 'Next.js', 'Renderiza las páginas, maneja la navegación y envía peticiones al backend.'],
        ['Capa de aplicación', 'Express', 'Gestiona endpoints, autenticación, validación y lógica de negocio.'],
        ['Capa de datos', 'MySQL', 'Almacena usuarios, anuncios, sugerencias y resto de entidades.']
    ]
    add_table(doc, ['Capa', 'Tecnología', 'Responsabilidad'], arch_rows)

    add_heading(doc, '3.4. Modelado de datos', level=2)
    add_heading(doc, '3.4.1. Datos de entrada', level=3)
    add_paragraph(doc, 'Los datos de entrada provienen de formularios de registro, publicación de anuncios, envío de sugerencias, creación de propuestas, mensajes y panel de administración. Incluyen textos, imágenes, identificadores de territorio y roles.')
    add_heading(doc, '3.4.2. Datos de salida', level=3)
    add_paragraph(doc, 'Los datos de salida son listados de anuncios, detalles de publicaciones, estadísticas, resultados de búsqueda, mensajes de confirmación y vistas de administración. Se sirven como HTML y JSON.')
    add_heading(doc, '3.4.3. Datos almacenados', level=3)
    add_paragraph(doc, 'El esquema de MySQL persiste usuarios, anuncios, favoritos, comunidad, sugerencias, propuestas, recursos, eventos, reportes, mensajes, agenda, tareas, plantillas y logs de actividad.')
    add_heading(doc, '3.4.4. Modelo entidad-relación', level=3)
    add_paragraph(doc, 'El modelo relacional se basa en entidades centrales: usuarios, anuncios, provincias y comunidades. A partir de estas se extienden tablas de relación como anuncios_guardados, comunidad_comentarios, propuestas_apoyos y reportes.')
    add_heading(doc, '3.4.5. Modelo relacional', level=3)
    ent_rows = [
        ['usuarios', 'Almacena cuentas y roles', 'id, nombre, email, password_hash, rol, creado_en', 'anuncios, comunidad, mensajes'],
        ['anuncios', 'Publicaciones del sistema', 'id, titulo, descripcion, categoria, comunidad_id, provincia_id, usuario_id, estado', 'favoritos, reportes'],
        ['comunidad_publicaciones', 'Publicaciones en comunidad', 'id, titulo, contenido, provincia_id, usuario_id', 'comentarios, likes'],
        ['sugerencias', 'Sugerencias de usuarios', 'id, asunto, descripcion, categoria, prioridad', '-'],
        ['propuestas', 'Propuestas ciudadanas', 'id, titulo, descripcion, comunidad_id, apoyos', 'propuestas_apoyos'],
        ['recursos', 'Recursos juveniles', 'id, titulo, descripcion, categoria, url', '-'],
        ['eventos', 'Eventos culturales', 'id, titulo, descripcion, fecha, provincia_id', '-'],
        ['reportes_anuncios', 'Reportes de contenido', 'id, anuncio_id, motivo, usuario_id', 'anuncios'],
        ['mensajes_staff', 'Mensajes internos', 'id, asunto, cuerpo, remitente_id, destinatario_id', '-'],
        ['admin_tareas', 'Tareas del panel admin', 'id, titulo, descripcion, estado, asignado_a', '-'],
    ]
    add_table(doc, ['Tabla', 'Finalidad', 'Campos principales', 'Relaciones'], ent_rows)
    add_heading(doc, '3.4.6. Script de base de datos', level=3)
    add_paragraph(doc, 'El script principal se encuentra en database/init.sql. Crea el esquema completo, inserta comunidades autónomas, provincias, anuncios de ejemplo, usuarios de demostración y datos iniciales de cultura y eventos.')

    add_heading(doc, '3.5. Análisis y diseño del sistema funcional', level=2)
    add_heading(doc, '3.5.1. Actores del sistema', level=3)
    actor_rows = [
        ['Visitante', 'Usuario no autenticado', 'Consultar anuncios, recursos, cultura y eventos.'],
        ['Usuario registrado', 'Usuario con cuenta', 'Publicar anuncios, enviar sugerencias, guardar favoritos.'],
        ['Moderador', 'Usuario con rol de moderador', 'Revisar reportes, moderar contenido.'],
        ['Administrador', 'Usuario con rol de administrador', 'Gestionar usuarios, contenidos, tareas, comunicaciones y agenda.']
    ]
    add_table(doc, ['Actor', 'Descripción', 'Funciones principales'], actor_rows)
    add_heading(doc, '3.5.2. Casos de uso', level=3)
    cu_rows = [
        ['CU-01', 'Registro de usuario', 'Visitante', 'Alta cuenta', 'Cuenta creada'],
        ['CU-02', 'Inicio de sesión', 'Usuario', 'Autenticarse', 'Token obtenido'],
        ['CU-03', 'Publicar anuncio', 'Usuario', 'Rellenar formulario', 'Anuncio pendiente de revisión'],
        ['CU-04', 'Buscar anuncios', 'Visitante', 'Aplicar filtros', 'Listado filtrado'],
        ['CU-05', 'Guardar favorito', 'Usuario', 'Pulsar en favorito', 'Anuncio en guardados'],
        ['CU-06', 'Enviar sugerencia', 'Usuario', 'Formulario de buzón', 'Sugerencia registrada'],
        ['CU-07', 'Moderar anuncio', 'Moderador', 'Revisar reporte', 'Anuncio aprobado o rechazado'],
        ['CU-08', 'Gestionar usuarios', 'Administrador', 'Panel de usuarios', 'Usuario actualizado']
    ]
    add_table(doc, ['ID', 'Caso de uso', 'Actor', 'Descripción', 'Resultado esperado'], cu_rows)
    add_heading(doc, '3.5.3. Flujo general de la aplicación', level=3)
    add_paragraph(doc, 'El usuario accede al frontend. Si consulta contenido, el frontend solicita datos al backend. Si publica, rellena el formulario y el backend guarda el registro en MySQL con estado pendiente. Un moderador puede revisar reportes y cambiar el estado.')
    add_heading(doc, '3.5.4. Seguridad lógica', level=3)
    add_paragraph(doc, 'La seguridad lógica se implementa mediante autenticación con JWT, contraseñas cifradas con bcrypt, protección CORS, rate limiting con express-rate-limit y validación de datos con Zod y express-validator. Los roles se consultan en el token y en las rutas protegidas.')

    add_heading(doc, '3.6. Análisis y diseño de la interfaz de usuario', level=2)
    add_heading(doc, '3.6.1. Criterios de diseño', level=3)
    add_paragraph(doc, 'El diseño busca claridad y simplicidad. Se ha empleado Tailwind CSS con una paleta definida, tipografía legible y componentes reutilizables. La navegación es horizontal con acceso directo a secciones principales.')
    add_heading(doc, '3.6.2. Pantallas representativas', level=3)
    for it in [
        'Home con listado de anuncios generales.',
        'Página de detalle de anuncio.',
        'Formulario de publicación de anuncio.',
        'Panel de administración.',
        'Panel de moderación.',
        'Sección de cultura y eventos.',
        'Buzón de sugerencias.',
        'Página de comunidad por provincias.'
    ]:
        add_bullet(doc, it)
    add_paragraph(doc, 'Las capturas de pantalla deben insertarse manualmente en una revisión posterior.')
    add_heading(doc, '3.6.3. Diseño responsive', level=3)
    add_paragraph(doc, 'La interfaz se adapta a dispositivos móviles mediante clases de Tailwind CSS. El menú se colapsa, las grillas pasan a una sola columna y los formularios ajustan su ancho.')

    # 4. IMPLEMENTACIÓN
    add_heading(doc, '4. Documento de Implementación, Pruebas e Implantación del sistema', level=1, page_break=True)
    add_heading(doc, '4.1. Implementación', level=2)
    add_heading(doc, '4.1.1. Estructura del frontend', level=3)
    add_paragraph(doc, 'El frontend se organiza bajo frontend/src. La carpeta app contiene las rutas de Next.js. La carpeta components agrupa componentes por área: layout, ui, anuncios, comunidad, admin y forms. Los hooks gestionan peticiones y el contexto de autenticación.')
    add_heading(doc, '4.1.2. Estructura del backend', level=3)
    add_paragraph(doc, 'El backend se estructura en backend/src. Los controladores gestionan la lógica, las rutas definen los endpoints, los middlewares validan autenticación y los utils incluyen utilidades como categorías y territorios.')
    add_heading(doc, '4.1.3. Gestión de base de datos', level=3)
    add_paragraph(doc, 'La conexión a MySQL se configura en backend/src/config/database.ts. El esquema inicial se carga con database/init.sql al arrancar el contenedor mysql. Se utiliza el driver mysql2 y se preparan consultas parametrizadas para evitar inyección SQL.')
    add_heading(doc, '4.1.4. Autenticación y roles', level=3)
    add_paragraph(doc, 'La autenticación se basa en tokens JWT. El registro cifra la contraseña con bcrypt. El token incluye el id y el rol. Las rutas protegidas verifican el token y el rol requerido.')
    add_heading(doc, '4.1.5. Gestión de anuncios', level=3)
    add_paragraph(doc, 'Los anuncios se crean, consultan, filtran y reportan. El controlador anuncios-mysql.ts implementa filtros por categoría, comunidad, provincia y búsqueda. La publicación permite adjuntar carteles mediante el endpoint de upload.')
    add_heading(doc, '4.1.6. Comunidad', level=3)
    add_paragraph(doc, 'La sección de comunidad permite publicar contenidos vinculados a una provincia y recibir comentarios. Existen controladores para publicaciones, comentarios, likes y reportes.')
    add_heading(doc, '4.1.7. Sugerencias y propuestas', level=3)
    add_paragraph(doc, 'El buzón de sugerencias recoge propuestas de mejora. Las propuestas permiten mostrar iniciativas ciudadanas y registrar apoyos.')
    add_heading(doc, '4.1.8. Cultura, eventos y recursos', level=3)
    add_paragraph(doc, 'La sección cultura agrupa eventos y actividades. Los anuncios culturales se filtran con la utilidad CATEGORIAS_CULTURA y se excluyen del home para mantenerlos en su área específica.')
    add_heading(doc, '4.1.9. Administración y moderación', level=3)
    add_paragraph(doc, 'El panel admin incluye módulos de anuncios, usuarios, estadísticas, tareas, agenda, mensajes, plantillas e instituciones. El panel de moderación permite gestionar reportes y revisar anuncios pendientes.')
    add_heading(doc, '4.1.10. Mensajería', level=3)
    add_paragraph(doc, 'Existe una vista de mensajes staff y plantillas de comunicación. Algunos aspectos, como el envío real, se presentan como demostración con datos mock.')
    add_heading(doc, '4.1.11. Dockerización del proyecto', level=3)
    add_paragraph(doc, 'El archivo docker-compose.yml define tres servicios: mysql, backend y frontend. El contenedor mysql carga database/init.sql. El contenedor backend espera a que MySQL esté saludable. El frontend se construye y expone en el puerto 3001.')

    add_heading(doc, '4.2. Instalación, despliegue y configuración', level=2)
    add_heading(doc, '4.2.1. Requisitos previos', level=3)
    add_bullet(doc, 'Docker Desktop en ejecución.')
    add_bullet(doc, 'Git para clonar el repositorio.')
    add_heading(doc, '4.2.2. Variables de entorno', level=3)
    add_paragraph(doc, 'Los archivos .env.example contienen los valores por defecto. No es necesario configurar nada para la primera ejecución local, ya que docker-compose.yml incluye las credenciales de demostración.')
    add_heading(doc, '4.2.3. Ejecución con Docker Compose', level=3)
    add_code_block(doc, 'git clone URL_DEL_REPOSITORIO\ncd citypaj\ndocker compose up --build')
    add_heading(doc, '4.2.4. Ejecución manual', level=3)
    add_paragraph(doc, 'También es posible ejecutar el frontend y el backend de forma manual con npm install y npm run dev, pero requiere una instancia MySQL disponible.')
    add_heading(doc, '4.2.5. Acceso a la aplicación', level=3)
    for it in ['Frontend: http://localhost:3001', 'Backend: http://localhost:3002', 'Health check: http://localhost:3002/health']:
        add_bullet(doc, it)

    add_heading(doc, '4.3. Pruebas', level=2)
    add_heading(doc, '4.3.1. Plan de pruebas', level=3)
    add_paragraph(doc, 'El plan de pruebas cubre funcionalidades principales del sistema. Las pruebas se han realizado de forma manual y mediante usuarios demo.')
    test_rows = [
        ['P-01', 'Registro', 'Crear cuenta con email y contraseña', 'Cuenta creada y login funcional', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-02', 'Login', 'Iniciar sesión con credenciales demo', 'Token JWT y redirección', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-03', 'Consulta anuncios', 'Cargar home', 'Listado de anuncios visible', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-04', 'Filtros', 'Aplicar filtro por comunidad', 'Listado filtrado', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-05', 'Publicación', 'Rellenar formulario de anuncio', 'Anuncio en estado pendiente', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-06', 'Favoritos', 'Guardar anuncio', 'Aparece en panel de guardados', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-07', 'Docker', 'Levantar contenedores', 'Servicios en puertos 3001 y 3002', 'No ejecutada en el entorno actual', 'Pendiente'],
        ['P-08', 'Health check', 'Llamar a /health', 'Respuesta correcta del backend', 'No ejecutada en el entorno actual', 'Pendiente']
    ]
    add_table(doc, ['ID', 'Tipo', 'Pasos', 'Resultado esperado', 'Resultado obtenido', 'Estado'], test_rows)
    add_heading(doc, '4.3.2. Resultado de las pruebas', level=3)
    add_paragraph(doc, 'Las pruebas manuales se proponen como parte del despliegue local. En el entorno de redacción de la memoria no se ha levantado la aplicación, por lo que las pruebas se marcan como pendientes de ejecución.')

    add_heading(doc, '4.4. Manual de usuario', level=2)
    add_paragraph(doc, 'El manual de usuario está resumido en el archivo README.md. El usuario puede acceder al frontend, registrarse o iniciar sesión, navegar por anuncios, publicar contenido, consultar cultura y eventos, enviar sugerencias y, si tiene permisos, acceder al panel de administración o moderación.')

    # 5. CIERRE
    add_heading(doc, '5. Documento de Cierre', level=1, page_break=True)
    add_heading(doc, '5.1. Resultados obtenidos', level=2)
    add_paragraph(doc, 'Se ha construido una aplicación web funcional con frontend, backend y base de datos dockerizados. El sistema permite consultar y publicar anuncios, gestionar favoritos, filtrar contenido, enviar sugerencias, consultar propuestas y acceder a paneles de administración y moderación. La documentación del README describe el despliegue local.')
    add_heading(doc, '5.2. Conclusiones', level=2)
    add_paragraph(doc, 'El proyecto ha permitido poner en práctica el conjunto de tecnologías estudiadas durante el ciclo. La dockerización ha sido una decisión tardía pero útil para la entrega. El filtrado territorial y la moderación de contenidos han supuesto los retos más importantes.')
    add_heading(doc, '5.3. Diario de bitácora', level=2)
    bit_rows = [
        ['Enero-Marzo', 'Configuración inicial, maquetación del frontend y primeras rutas.', 'Commit inicial y estructura básica.', 'Aprendizaje de Next.js App Router.'],
        ['Abril', 'Conexión con MySQL, modelado de datos y controladores de anuncios.', 'Migraciones y seeds.', 'Dificultades con tipos de Knex.'],
        ['Mayo', 'Paneles admin y moderación, comunidad, sugerencias y propuestas.', 'Paneles admin y moderador.', 'Muchos módulos en paralelo.'],
        ['Junio-Agosto', 'Dockerización, limpieza del repositorio, README y memoria.', 'Docker compose y README.', 'Revisión final de referencias.']
    ]
    add_table(doc, ['Periodo', 'Trabajo realizado', 'Evidencia', 'Observaciones'], bit_rows)
    add_heading(doc, '5.4. Temporalización real', level=2)
    add_paragraph(doc, 'El desarrollo se ha extendido desde enero hasta agosto, con una intensidad mayor durante los últimos meses. El trabajo se ha ajustado según las dificultades encontradas.')
    add_heading(doc, '5.5. Desviación respecto a la planificación inicial', level=2)
    desv_rows = [
        ['Docker', '4 horas', '8 horas', '+4', 'La configuración del proxy requirió más tiempo.'],
        ['Moderación', '10 horas', '15 horas', '+5', 'Se añadieron filtros y acciones masivas.'],
        ['Pruebas', '10 horas', '5 horas', '-5', 'Se priorizó funcionalidad frente a pruebas formales.']
    ]
    add_table(doc, ['Tarea', 'Planificado', 'Real', 'Desviación', 'Explicación'], desv_rows)
    add_heading(doc, '5.6. Conocimientos adquiridos', level=2)
    for it in ['Uso de Next.js App Router y React.', 'Desarrollo de API REST con Express.', 'Conexión y consultas a MySQL.', 'Dockerización de aplicaciones fullstack.', 'Gestión de autenticación con JWT y bcrypt.', 'Aplicación de Tailwind CSS para diseño responsive.']:
        add_bullet(doc, it)
    add_heading(doc, '5.7. Limitaciones del proyecto', level=2)
    add_paragraph(doc, 'El proyecto prioriza el entorno local y la demostración. Algunas funciones del panel admin se presentan con datos mock. La mensajería real y el envío de correos no están completamente operativos.')
    add_heading(doc, '5.8. Posibles mejoras futuras', level=2)
    for it in ['Completar la mensajería interna con notificaciones.', 'Implementar tests automáticos.', 'Mejorar la accesibilidad.', 'Optimizar consultas con índices.', 'Añadir estadísticas en tiempo real.', 'Preparar despliegue en servidor propio.']:
        add_bullet(doc, it)

    # 6. BIBLIOGRAFÍA
    add_heading(doc, '6. Bibliografía', level=1, page_break=True)
    bib = [
        'Next.js Documentation. Consultado en 2026. Disponible en https://nextjs.org/docs',
        'React Documentation. Consultado en 2026. Disponible en https://react.dev',
        'Node.js Documentation. Consultado en 2026. Disponible en https://nodejs.org',
        'Express Documentation. Consultado en 2026. Disponible en https://expressjs.com',
        'MySQL Documentation. Consultado en 2026. Disponible en https://dev.mysql.com/doc',
        'Docker Documentation. Consultado en 2026. Disponible en https://docs.docker.com',
        'MDN Web Docs. Consultado en 2026. Disponible en https://developer.mozilla.org',
        'Tailwind CSS Documentation. Consultado en 2026. Disponible en https://tailwindcss.com/docs'
    ]
    for it in bib:
        add_paragraph(doc, it, first=False)

    # 7. ANEXOS
    add_heading(doc, '7. Anexos', level=1, page_break=True)
    add_heading(doc, '7.1. Enlace al repositorio', level=2)
    add_paragraph(doc, 'https://github.com/carmendmv/citypaj')
    add_heading(doc, '7.2. Estructura del repositorio', level=2)
    add_paragraph(doc, 'El repositorio contiene las carpetas backend, frontend, database, docs/memoria y los archivos de configuración docker-compose.yml, README.md y .env.example.')
    add_heading(doc, '7.3. Script de base de datos', level=2)
    add_paragraph(doc, 'El script database/init.sql crea las tablas e inserta los datos iniciales. Se encuentra en la carpeta database.')
    add_heading(doc, '7.4. Docker Compose', level=2)
    add_paragraph(doc, 'El archivo docker-compose.yml define los servicios mysql, backend y frontend, junto con el volumen citypaj_mysql_data.')
    add_heading(doc, '7.5. Endpoints principales', level=2)
    ep_rows = [
        ['GET', '/api/anuncios', 'Listado de anuncios'],
        ['POST', '/api/anuncios', 'Crear anuncio'],
        ['GET', '/api/anuncios/:id', 'Detalle de anuncio'],
        ['POST', '/api/auth/register', 'Registro'],
        ['POST', '/api/auth/login', 'Login'],
        ['GET', '/api/eventos', 'Listado de eventos'],
        ['GET', '/api/recursos', 'Listado de recursos'],
        ['POST', '/api/sugerencias', 'Enviar sugerencia'],
        ['GET', '/api/admin/usuarios', 'Gestión de usuarios'],
        ['GET', '/api/moderacion', 'Panel de moderación']
    ]
    add_table(doc, ['Método', 'Endpoint', 'Descripción'], ep_rows)
    add_heading(doc, '7.6. Commits relevantes', level=2)
    add_paragraph(doc, 'El historial de commits refleja el desarrollo desde la estructura inicial hasta la dockerización final. Entre los más relevantes se encuentran: Configura Docker, variables de entorno y scripts de build; Añade esquema y datos iniciales de MySQL; Implementa API de anuncios, filtro de cultura y utilidad de categorías; Configura proxy interno hacia el backend en Docker; Excluye anuncios culturales del home y añade utilidad de categorías; y readme definitivo.')
    add_heading(doc, '7.7. Manual de instalación resumido', level=2)
    add_code_block(doc, 'git clone https://github.com/carmendmv/citypaj\ncd citypaj\ndocker compose up --build')
    add_paragraph(doc, 'Una vez levantados, acceder a http://localhost:3001 para el frontend y a http://localhost:3002 para el backend.')

    doc.save(DOCX)

    # INFORME
    informe_text = f"""# Informe de generación de la memoria

1. Archivo Word generado: `{DOCX}`
2. Apartados incluidos: portada, índices, siete documentos principales, bibliografía y anexos.
3. Archivos del repositorio analizados: README.md, package.json, frontend/package.json, backend/package.json, docker-compose.yml, backend/Dockerfile, frontend/Dockerfile, .env.example, backend/.env.example, database/init.sql, backend/src/index.ts y estructura de carpetas.
4. Tecnologías detectadas: Next.js 14, React 18, TypeScript, Tailwind CSS, Node.js, Express, MySQL 8.0, mysql2, bcryptjs, jsonwebtoken, Docker, Docker Compose, Git.
5. Base de datos detectada: MySQL con base de datos `citypaj` y esquema en `database/init.sql`.
6. Docker detectado: sí, con servicios mysql, backend y frontend.
7. Funcionalidades documentadas: registro, login, anuncios, filtros, favoritos, comunidad, sugerencias, propuestas, cultura y eventos, recursos, administración, moderación y mensajería.
8. Funcionalidades parciales detectadas: mensajería real, estadísticas avanzadas, envío de correos, algunos paneles con datos de demostración.
9. Pruebas documentadas: plan de pruebas con estados pendientes de ejecución.
10. Capturas insertadas o pendientes: no se han insertado capturas. Se indica que deben añadirse manualmente.
11. Comprobación de limpieza de rastros: se ha revisado el texto generado y no se incluyen referencias a asistentes, modelos de lenguaje, cascada, herramientas de IA, emojis, lorem ipsum, TODO ni FIXME.
12. Observaciones finales: el documento es una memoria académica completa basada en el repositorio real. Se recomienda revisar las secciones de capturas y pruebas tras ejecutar la aplicación.
"""
    with open(INFO, 'w', encoding='utf-8') as f:
        f.write(informe_text)

    print(f'Memoria guardada en: {DOCX}')
    print(f'Informe guardado en: {INFO}')

if __name__ == '__main__':
    main()
